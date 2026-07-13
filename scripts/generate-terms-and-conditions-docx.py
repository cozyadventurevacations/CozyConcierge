from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_ROOT / "public" / "legal"
OUTPUT_PATH = OUTPUT_DIR / "cozy-adventure-vacations-terms-and-conditions.docx"

BRAND_BLUE = RGBColor(18, 63, 91)
ACCENT_BLUE = RGBColor(44, 95, 138)
MUTED = RGBColor(102, 112, 133)


SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Acceptance of These Terms",
        [
            "These Terms and Conditions govern travel planning, booking, and related services provided by Cozy Adventure Vacations LLC (\"Cozy Adventure Vacations,\" \"we,\" \"our,\" or \"us\"). By requesting a quote, approving an itinerary, submitting payment, signing an authorization, or otherwise booking travel through us, each client, traveler, and authorized payer agrees to these Terms.",
            "If a trip includes more than one traveler, the person who approves the booking represents that they are authorized to accept these Terms on behalf of all travelers included in the booking and is responsible for sharing all applicable trip information with those travelers.",
        ],
    ),
    (
        "2. Our Role as Travel Advisor",
        [
            "Cozy Adventure Vacations provides travel planning, booking support, itinerary coordination, and concierge-style assistance. We arrange travel through independent third-party suppliers, including cruise lines, tour operators, resorts, hotels, airlines, transportation companies, excursion providers, insurance providers, and other vendors.",
            "We are not the owner, operator, or controlling party for supplier services. Supplier schedules, prices, inventory, cancellation terms, refund rules, loyalty benefits, accessibility accommodations, safety standards, and service delivery are controlled by the applicable supplier.",
        ],
    ),
    (
        "3. Supplier Terms and Conditions",
        [
            "All travel components are subject to the terms, conditions, rules, tariffs, contracts of carriage, cancellation policies, health and safety rules, documentation requirements, and refund policies of the applicable supplier. Supplier terms may be more restrictive than these Terms and may change without notice.",
            "Clients are responsible for reviewing supplier terms before approving a booking. If supplier terms conflict with these Terms as to a supplier-provided service, the supplier terms control for that service. These Terms continue to govern the relationship between the client and Cozy Adventure Vacations.",
        ],
    ),
    (
        "4. Quotes, Pricing, and Availability",
        [
            "All quotes are subject to availability and are not guaranteed until the booking is confirmed by the supplier and required payments have been received and processed. Prices, taxes, fees, exchange rates, fuel surcharges, resort fees, port charges, gratuities, baggage charges, seat fees, and other costs may change before confirmation.",
            "A quote may include estimated costs based on information available at the time it is prepared. Final pricing is controlled by the supplier and may differ from the estimate if availability, currency rates, taxes, fees, or supplier policies change.",
        ],
    ),
    (
        "5. Payments, Deposits, and Authorization",
        [
            "Deposits, interim payments, and final payments are due by the deadlines provided for the booking. Missing a payment deadline may result in cancellation, loss of space, repricing, penalties, or forfeiture of amounts paid, according to supplier terms.",
            "The client is responsible for ensuring that all payments are submitted on time. Cozy Adventure Vacations may send reminders and provide secure payment links, but reminders are a courtesy and do not shift payment responsibility away from the client.",
            "Payments may be processed directly by suppliers, payment processors, or other authorized vendors. Clients should never send credit card numbers, passport scans, passwords, or other sensitive information by ordinary email or text. Secure portal links or approved payment authorization processes should be used for sensitive information.",
            "If a separate planning fee, service fee, cancellation fee, change fee, or administrative fee is disclosed for a booking, that fee is payable as stated and may be separate from supplier charges. Cozy Adventure Vacations fees are generally non-refundable once planning, booking, or service work has begun unless otherwise stated in writing.",
        ],
    ),
    (
        "6. Changes, Cancellations, and Refunds",
        [
            "Change and cancellation rules vary by supplier and may include non-refundable deposits, penalties, fare differences, repricing, or loss of promotional benefits. Clients should review all applicable terms before authorizing changes or cancellations.",
            "Cozy Adventure Vacations will assist with supplier communications and refund requests when appropriate, but we do not control supplier decisions, refund eligibility, refund method, processing time, credit validity, or the availability of alternative travel arrangements.",
            "Where an airline, supplier, or other provider is legally required to issue a refund, Cozy Adventure Vacations will assist the client in seeking the applicable refund or credit from the responsible provider, but we are not responsible for funds held or delayed by third parties.",
            "Unauthorized chargebacks or payment disputes may be contested if services were requested, booked, or provided. Clients should contact Cozy Adventure Vacations promptly so we can help address payment, refund, or supplier concerns before initiating a payment dispute.",
        ],
    ),
    (
        "7. Travel Documents and Entry Requirements",
        [
            "Travelers are solely responsible for obtaining, carrying, and presenting all required travel documents, including passports, visas, government identification, entry forms, vaccination records, parental or guardian consent documents, cruise boarding documents, and any other documents required by governments or suppliers.",
            "Entry requirements can change at any time. Travelers should review official government sources and destination-specific advisories before departure and again close to travel. Many destinations and suppliers require passports to be valid for a period beyond the travel dates and may require blank passport pages.",
            "Cozy Adventure Vacations may provide general reminders or links as a courtesy, but we do not issue legal, immigration, medical, or government-entry advice and are not responsible for denied boarding, denied entry, quarantine, penalties, delays, or additional costs caused by missing, expired, incorrect, or insufficient documentation.",
        ],
    ),
    (
        "8. Travel Insurance",
        [
            "Cozy Adventure Vacations strongly recommends that every traveler purchase comprehensive travel insurance or travel protection that fits the trip and traveler needs, including coverage for cancellation, interruption, medical emergencies, evacuation, supplier default, baggage loss or delay, travel delay, and other unexpected events.",
            "Insurance benefits, exclusions, pre-existing condition rules, purchase deadlines, claim requirements, and coverage limits are determined by the insurer and policy terms. Travelers are responsible for reviewing policy documents and deciding whether coverage is adequate.",
            "If a traveler declines travel insurance or fails to purchase coverage in time, the traveler accepts the risk of uncovered losses and releases Cozy Adventure Vacations from responsibility for losses that could have been covered by insurance or travel protection.",
        ],
    ),
    (
        "9. Air Travel, Schedule Changes, and Independent Arrangements",
        [
            "Airline schedules, seat assignments, aircraft, routes, baggage rules, check-in deadlines, and cancellation or delay handling are controlled by the airline. Travelers are responsible for monitoring flight schedules, checking in on time, meeting baggage rules, and arriving at the airport with sufficient time.",
            "If travelers book their own airfare, hotels, transfers, excursions, or other arrangements outside the itinerary booked through Cozy Adventure Vacations, those independent arrangements remain the travelers' responsibility. We are not liable for conflicts, missed connections, cancellation penalties, or supplier issues related to independently booked components.",
        ],
    ),
    (
        "10. Special Requests and Accessibility Needs",
        [
            "Clients should notify Cozy Adventure Vacations as early as possible of mobility needs, medical considerations, dietary restrictions, bedding preferences, connecting rooms, celebrations, accessibility requests, or other special requests. We will communicate requests to suppliers when feasible.",
            "Special requests are not guaranteed unless confirmed directly by the supplier in writing. Suppliers may require additional documentation, advance notice, or fees. Travelers are responsible for confirming that travel arrangements are suitable for their personal, medical, mobility, dietary, or accessibility needs.",
        ],
    ),
    (
        "11. Health, Safety, and Traveler Conduct",
        [
            "Travel involves inherent risks, including illness, injury, weather events, political unrest, supplier disruptions, schedule changes, transportation delays, and other circumstances beyond our control. Travelers are responsible for evaluating personal fitness to travel and for following all supplier, government, safety, and conduct rules.",
            "Suppliers may refuse service, remove travelers, or impose penalties for disruptive conduct, failure to comply with rules, failure to meet health or documentation requirements, or conduct that affects the safety or experience of others. Any resulting costs or losses are the responsibility of the traveler.",
        ],
    ),
    (
        "12. Force Majeure and Events Beyond Our Control",
        [
            "Cozy Adventure Vacations is not liable for delay, cancellation, loss, injury, expense, inconvenience, or failure to perform caused by circumstances beyond our reasonable control, including weather, natural disasters, pandemics, epidemics, government action, war, terrorism, civil unrest, labor actions, supplier insolvency, mechanical failures, technology outages, airport or port closures, schedule changes, or other force majeure events.",
        ],
    ),
    (
        "13. Limitation of Responsibility",
        [
            "To the fullest extent permitted by law, Cozy Adventure Vacations is not responsible for acts, errors, omissions, representations, warranties, breaches, negligence, bankruptcy, insolvency, cancellations, schedule changes, delays, injuries, death, property damage, emotional distress, inconvenience, or other losses caused by third-party suppliers or circumstances outside our control.",
            "Our role is limited to advisor and booking support services. Any claim relating to supplier-provided transportation, lodging, tours, cruises, excursions, insurance, or other travel services must be directed to the responsible supplier, insurer, or provider.",
        ],
    ),
    (
        "14. Client Information, Privacy, and Electronic Communications",
        [
            "Clients are responsible for providing accurate names, dates of birth, contact details, passport information, traveler preferences, payment information, and other trip information. Names must match government-issued identification and travel documents.",
            "Cozy Adventure Vacations may communicate by email, text message, phone, secure portal, electronic signature, or other electronic means. Clients consent to electronic communications related to quotes, bookings, documents, payment reminders, schedule changes, insurance decisions, and trip servicing.",
            "We take reasonable steps to protect client information, but ordinary email and text messaging are not secure channels for sensitive information. Clients should use secure upload, portal, or payment links when submitting sensitive documents or payment information.",
        ],
    ),
    (
        "15. Intellectual Property and Portal Use",
        [
            "Itineraries, proposals, planning materials, portal content, checklists, recommendations, and other materials prepared by Cozy Adventure Vacations are for the client's personal trip-planning use and may not be copied, resold, republished, or used for commercial purposes without written permission.",
            "Clients are responsible for maintaining the confidentiality of portal login credentials and for promptly reporting suspected unauthorized access.",
        ],
    ),
    (
        "16. Governing Law and Dispute Resolution",
        [
            "These Terms are governed by the laws of the State of Illinois, without regard to conflict-of-law principles, unless another jurisdiction is required by applicable law.",
            "Clients agree to first contact Cozy Adventure Vacations in writing and allow a reasonable opportunity to resolve any concern informally. If a dispute cannot be resolved informally, the parties agree that any legal action against Cozy Adventure Vacations shall be brought in a court of competent jurisdiction in Illinois, unless prohibited by applicable law.",
        ],
    ),
    (
        "17. Updates to These Terms",
        [
            "Cozy Adventure Vacations may update these Terms from time to time. The version posted in the client portal or provided with booking documents applies to services requested or bookings made after the effective date shown on that version, unless otherwise stated in writing.",
        ],
    ),
    (
        "18. Contact Information",
        [
            "Cozy Adventure Vacations LLC",
            "Email: CozyAdventureVacations@gmail.com",
            "Phone: (630) 219-1904",
            "Website: www.CozyAdventureVacations.com",
        ],
    ),
]


def set_run_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, field: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_BLUE, 16, 8),
        ("Heading 2", 13, ACCENT_BLUE, 12, 6),
        ("Heading 3", 12, BRAND_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Cozy Adventure Vacations LLC - Terms and Conditions - Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "PAGE")


def build_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Cozy Adventure Vacations LLC")
    set_run_font(run, size=22, color=BRAND_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Terms and Conditions")
    set_run_font(run, size=16, color=ACCENT_BLUE, bold=True)

    effective = doc.add_paragraph()
    effective.alignment = WD_ALIGN_PARAGRAPH.CENTER
    effective.paragraph_format.space_after = Pt(18)
    run = effective.add_run("Effective Date: July 13, 2026")
    set_run_font(run, size=10, color=MUTED, italic=True)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "Please read these Terms and Conditions carefully. They explain the responsibilities of Cozy Adventure Vacations LLC, third-party travel suppliers, and each traveler when travel services are requested, quoted, booked, changed, or cancelled."
    )
    set_run_font(run)

    for heading, paragraphs in SECTIONS:
        doc.add_paragraph(heading, style="Heading 1")
        for text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.10
            run = p.add_run(text)
            set_run_font(run)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
